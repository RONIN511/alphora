"""
Word 文档查看器 - 处理 .docx/.doc 文件
"""
import re
from typing import Optional, List, Dict, Any, Tuple

from ..utils.common import get_file_info, truncate_text


class DocumentViewer:
    """Word 文档查看器"""
    
    SUPPORTED_EXTENSIONS = {'.docx', '.doc'}
    
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.file_info = get_file_info(file_path)
        self.ext = self.file_info['extension']
        
    def view(
        self,
        purpose: str = "preview",
        keyword: Optional[str] = None,
        max_lines: int = 100,
        page_number: Optional[int] = None,
    ) -> str:
        """查看 Word 文档内容"""
        purpose, warnings = self._infer_params(purpose, keyword)
        
        try:
            from docx import Document
        except ImportError:
            return "❌ 需要安装 python-docx: pip install python-docx"
        
        try:
            doc = Document(self.file_path)
        except Exception as e:
            return f"❌ 无法打开文档: {e}"
        
        paragraphs = self._extract_paragraphs(doc)
        tables = self._extract_tables(doc)
        
        if purpose == "structure":
            return self._get_structure(paragraphs, tables, warnings)
        elif purpose == "search":
            return self._search(paragraphs, tables, keyword, max_lines, warnings)
        else:
            return self._preview(paragraphs, tables, max_lines, warnings)
    
    def _infer_params(self, purpose: str, keyword: Optional[str]) -> Tuple[str, List[str]]:
        warnings = []
        if keyword and purpose != "search":
            warnings.append(f"⚠️ 检测到 keyword='{keyword}'，已自动切换为 search 模式")
            purpose = "search"
        if purpose == "search" and not keyword:
            purpose = "preview"
        return purpose, warnings
    
    def _extract_paragraphs(self, doc) -> List[Dict[str, Any]]:
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                style = para.style.name if para.style else "Normal"
                paragraphs.append({
                    'text': text,
                    'style': style,
                    'is_heading': 'Heading' in style,
                })
        return paragraphs
    
    def _extract_tables(self, doc) -> List[List[List[str]]]:
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        return tables
    
    def _format_header(self, paragraphs, tables, warnings) -> str:
        lines = [
            f"📄 文件: {self.file_info['name']}",
            f"📦 大小: {self.file_info['size_human']}",
            f"📋 段落数: {len(paragraphs)} | 表格数: {len(tables)}",
        ]
        if warnings:
            lines.extend([""] + warnings)
        return '\n'.join(lines)
    
    def _get_structure(self, paragraphs, tables, warnings) -> str:
        lines = [self._format_header(paragraphs, tables, warnings), "", "【文档结构】"]
        headings = [p for p in paragraphs if p['is_heading']]
        if headings:
            for h in headings[:30]:
                lines.append(f"• {truncate_text(h['text'], 60)}")
        else:
            lines.append("  (文档没有标题结构)")
        return '\n'.join(lines)
    
    def _preview(self, paragraphs, tables, max_lines, warnings) -> str:
        lines = [self._format_header(paragraphs, tables, warnings), "", "【内容预览】"]
        for p in paragraphs[:max_lines]:
            lines.append(p['text'])
        if len(paragraphs) > max_lines:
            lines.append(f"\n... 还有 {len(paragraphs) - max_lines} 个段落")
        return '\n'.join(lines)
    
    def _search(self, paragraphs, tables, keyword, max_lines, warnings) -> str:
        results = []
        keyword_lower = keyword.lower()
        
        for i, p in enumerate(paragraphs, 1):
            if keyword_lower in p['text'].lower():
                results.append({'type': 'paragraph', 'location': f"段落{i}", 'content': p['text'][:100]})
        
        lines = [
            self._format_header(paragraphs, tables, warnings), "",
            f"🔍 搜索: '{keyword}'",
            f"📋 找到 {len(results)} 处匹配", ""
        ]
        
        if results:
            for r in results[:max_lines]:
                lines.append(f"[{r['location']}] {r['content']}")
        else:
            lines.append(f"未找到包含 '{keyword}' 的内容")
        
        return '\n'.join(lines)
